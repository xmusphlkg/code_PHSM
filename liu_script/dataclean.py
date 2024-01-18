import os
from html import unescape
from urllib.parse import unquote
import numpy as np
import xlrd
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import requests


def remove_chinese(title):
    title_type = title.split('.')[-1]
    title = ''.join([x for x in title if x.isdigit() or x == '-'])
    return title, title_type


def filetype(title):
    title_type = title.split('.')[-1]
    return title_type


def remove_space(title):
    title = title.replace(' ', '')
    return title


def process_files_combined(directory):
    data_list = []
    datatype_list = (np.floating,np.integer,int, float)

    files = os.listdir(directory)

    for file in files:
        file_path = os.path.join(directory, file)
        file_size = os.path.getsize(file_path)

        if file_size > 100:
            file_type = filetype(file)
            if file_type == 'doc':
                pass
            if file_type == 'docx':
                text_content, table_content = read_docx(file_path)
                df = pd.DataFrame(table_content)
                for i in range(len(df)):
                    try:
                        df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                    except:
                        pass
                    if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                            df.iloc[i][2], datatype_list):
                        data_list.append(
                            [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
            elif file_type == 'xls':
                sheet = xlrd.open_workbook(file_path).sheet_by_index(0)
                num_rows = sheet.nrows
                num_cols = sheet.ncols
                data = []

                for row_index in range(num_rows):
                    row_data = []

                    for col_index in range(num_cols):
                        cell_value = sheet.cell_value(row_index, col_index)
                        row_data.append(cell_value)

                    data.append(row_data)

                df = pd.DataFrame(data)

                for i in range(len(df)):
                    try:
                        df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                    except:
                        pass
                    if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                            df.iloc[i][2], datatype_list):
                        data_list.append(
                            [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
            elif file_type == 'csv':
                df = pd.read_csv(file_path, encoding='gbk')
                for i in range(len(df)):
                    if str(df.iloc[0, 0]) == '0':
                        try:
                            df.iloc[i, 2] = float(df.iloc[i, 2])
                            df.iloc[i, 3] = float(df.iloc[i, 3])
                        except:
                            pass
                        if isinstance(df.iloc[i][1], str) and isinstance(df.iloc[i][2], datatype_list) and isinstance(
                                df.iloc[i][3], datatype_list):
                            data_list.append(
                                [remove_space(df.iloc[i][1]), df.iloc[i][2], df.iloc[i][3], file.split('.')[0]])
                    else:
                        try:
                            df.iloc[i, 1] = float(df.iloc[i, 1])
                            df.iloc[i, 2] = float(df.iloc[i, 2])
                        except:
                            pass
                        try:
                            if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i,1],datatype_list) and isinstance(
                                    df.iloc[i][2], datatype_list):
                                data_list.append(
                                    [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
                        except:
                            pass
            elif file_type == 'xlsx':
                sheet = pd.read_excel(file_path)
                df = sheet

                for i in range(len(df)):
                    try:
                        df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                    except:
                        pass
                    if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                            df.iloc[i][2], datatype_list):
                        data_list.append(
                            [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])

    result_df = pd.DataFrame(data_list, columns=['疾病病种', '发病数', '死亡数', 'date'])
    return result_df


def read_docx(file_path):
    doc = Document(file_path)
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(row_content)
    return text_content, table_content


def get_year_month(title):
    year = title.split('布')[1].split('年')[0]
    month = title.split('年')[1].split('月')[0]
    return year, month


def download_and_parse_data(df_row, headers):
    url = df_row['链接']
    sign = 0

    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gbk')
    links_with_blank_target = soup.find_all('a', {'target': '_blank'})
    for link_tag in links_with_blank_target:
        link = link_tag.get('href')
        if link.endswith('docx'):
            sign = 1
            url_link = unescape(str(url) + link)
            response = requests.get(url_link, headers=headers)
            if response.status_code == 200:
                with open(f'./data/province/anhui/{df_row["年份"]}-{df_row["月份"]}.docx', 'wb') as f:
                    f.write(response.content)
            else:
                print(f'{df_row["年份"]}-{df_row["月份"]}存在，下载失败，但是200')

    try:
        table = soup.find('tbody')
        rows = table.find_all('tr')
        table_data = []
        for row in rows:
            cells = row.find_all(['td', 'th'])
            row_data = [cell.text.strip() for cell in cells]
            table_data.append(row_data)

        table_df = pd.DataFrame(table_data)
        table_df.replace('\xa0', '', regex=True, inplace=True)
        table_df.to_csv(f'./data/province/anhui/{df_row["年份"]}-{df_row["月份"]}.csv', encoding='gbk')
        sign = 1
    except Exception as e:
        print(f"Error processing table data: {e}")

    if sign == 0:
        print(f'{df_row["年份"]}-{df_row["月份"]}不存在传染病信息')


def scrape_and_save_table(soup, df, i):
    table = soup.find('tbody')
    rows = table.find_all('tr')
    table_data = []

    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_data = [cell.text.strip() for cell in cells]
        table_data.append(row_data)

    table_df = pd.DataFrame(table_data)
    table_df.replace('\xa0', '', regex=True, inplace=True)

    file_path = f'./data/province/anhui/{df["年份"].iloc[i]}-{df["月份"].iloc[i]}.csv'
    table_df.to_csv(file_path, encoding='gbk')


# 使用方式：
# scrape_and_save_table(soup, df, i)

def update_url_column(df_csv_path, url_csv_path):
    url_df = pd.read_csv(url_csv_path, encoding='gbk')
    df = pd.read_csv(df_csv_path, encoding='gbk')
    df['date_1'] = pd.to_datetime(df['date'])
    df['year'] = df['date_1'].dt.year.astype(str)
    df['month'] = df['date_1'].dt.month.astype(str)
    df['url'] = None

    for i in range(len(url_df)):
        url = unquote(url_df.loc[i, '链接'])
        url_year = url_df['年份'].iloc[i]
        url_month = url_df['月份'].iloc[i]

        for j in range(len(df)):
            if str(df['year'].iloc[j]) == str(url_year) and str(df['month'].iloc[j]) == str(url_month):
                df['url'].iloc[j] = url
    df.drop('date_1', axis=1, inplace=True)
    df.to_csv(df_csv_path, encoding='gbk', index=False)
