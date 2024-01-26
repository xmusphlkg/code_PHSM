import requests
from bs4 import BeautifulSoup
import re
from html import unescape
import os
import pandas as pd
from docx import Document
from dataclean import remove_space
from urllib.parse import unquote,quote

files = os.listdir('./data/province/jiangsu')
# Function to get year and month from title
def get_year_month(title):
    year = title.split('布')[1].split('年')[0]
    month = title.split('年')[1].split('月')[0]
    return year, month

# Function to read .docx files
def read_docx(file_path):
    doc = Document(file_path)

    # Read the text paragraph by paragraph
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)

    # Read content table by table
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(row_content)

    return text_content, table_content

# Get data from report
data = []
for i in range(1, 7):
    print(i)
    page=i
    headers = {
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
        'Accept-Encoding': 'gzip, deflate, br',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'Connection': 'keep-alive',
        'Host': 'www.jiangsu.gov.cn',
        'Referer': 'https://www.jiangsu.gov.cn/jrobot/search.do?webid=31&analyzeType=1&pg=10&p='+str(page+1)+'&tpl=2&category=&q='+quote('报告 传染病')+'&pos=title&od=1&date=&date=',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 Edg/122.0.0.0',
        'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Microsoft Edge";v="122"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
        'Content-Type': 'text/html;charset=UTF-8'
    }
    cookies = {
        'user_sid': '85fbbaa838e747ee9c081c5e5aa42cf8',
        'JSESSIONID': '9E24E70D857AB2A30EF53A3B2E9EC2C4',
        '__jsluid_s': '09431aa0db8838ce46bca374f82e9978',
    }
    url = "https://www.jiangsu.gov.cn/jrobot/search.do"
    params = {
        "webid": "31",
        "analyzeType": "1",
        "pg": "10",
        "p": str(page),
        "tpl": "2",
        "category": "",
        "q": "报告 传染病",
        "pos": "title",
        "od": "1",
        "date": "",
    }

    response = requests.get(url, headers=headers, params=params,cookies=cookies)
    soup=BeautifulSoup(response.text,'html.parser')
    links_with_blank_target = soup.find_all('a', {'target': '_blank'})
    url_origin = 'https://www.jiangsu.gov.cn'
    for link in links_with_blank_target:
        chinese_text = link.get_text(strip=True)
        url_link = str(url_origin)+link.get('href')
        data.append([chinese_text, url_link])
# Save data to csv file
df = pd.DataFrame(data, columns=['中文解释', '链接'])
df = df[df['中文解释'].str.contains('月全省')]
df['年份'],df['月份'] = zip(*df['中文解释'].apply(lambda x: get_year_month(x)))
df.to_csv('./data/province/jiangsu/jiangsu_url.csv', index=False, encoding='gbk')

# Download report
df['链接']=df['链接'].str.extract(r'url=(.*?)(?:&q=|$)')
#构造请求头
cookies = {
    'user_sid': '85fbbaa838e747ee9c081c5e5aa42cf8',
    'JSESSIONID': '9E24E70D857AB2A30EF53A3B2E9EC2C4',
    '__jsluid_s': '09431aa0db8838ce46bca374f82e9978',
}
headers = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
    "Accept-Encoding": "gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Host": "wjw.jiangsu.gov.cn",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Sec-Fetch-User": "?1",
    "Upgrade-Insecure-Requests": "1",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 Edg/122.0.0.0",
    "sec-ch-ua": '"Chromium";v="122", "Not(A:Brand";v="24", "Microsoft Edge";v="122"',
    "sec-ch-ua-mobile": "?0",
    "sec-ch-ua-platform": '"Windows"'
}

for i in range(len(df)):
    url = unquote(df['链接'].iloc[i])
    response = requests.get(url, headers=headers, cookies=cookies)
    soup=BeautifulSoup(response.text,'html.parser')
    links_with_blank_target = soup.find_all('a')
    link=re.compile(r'/module/download/.*?.docx').findall(str(links_with_blank_target))
    url_link = unescape(str(url_origin) + link[0])
    response=requests.get(url_link, headers=headers, cookies=cookies)
    if response.status_code == 200:
        with open(f'./data/province/jiangsu/{df["年份"].iloc[i]}-{df["月份"].iloc[i]}.docx', 'wb') as f:
            f.write(response.content)
data_list = []
for file in files:
    text_content, table_content = read_docx('./data/province/jiangsu/'+file)
    df = pd.DataFrame(table_content)
    for i in range(len(df)):
        try:
            df.iloc[i][1],df.iloc[i][2] = float(df.iloc[i][1]),float(df.iloc[i][2])
        except:
            pass
        if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], (int, float)) and isinstance(
                df.iloc[i][2], (int, float)):
            data_list.append([remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
result_df = pd.DataFrame(data_list, columns=['疾病病种', '发病数', '死亡数', 'date'])
result_df.to_csv('./data/province/jiangsu/jiangsu.csv', index=False, encoding='gbk')

url_df=pd.read_csv('./data/province/jiangsu/jiangsu_url.csv',encoding='gbk')
jiangsu=pd.read_csv('./data/province/jiangsu/jiangsu.csv',encoding='gbk')
jiangsu['date_1'] = pd.to_datetime(jiangsu['date'])
jiangsu['year'] = jiangsu['date_1'].dt.year.astype(str)
jiangsu['month']=jiangsu['date_1'].dt.month.astype(str)
jiangsu['url']=None

for i in range(len(url_df)):
    url=url_df.loc[i,'链接']
    url=unquote(url)
    url_year=url_df['年份'].iloc[i]
    url_month=url_df['月份'].iloc[i]
    for j in range(len(jiangsu)):
        if str(jiangsu['year'].iloc[j])==str(url_year) and str(jiangsu['month'].iloc[j])==str(url_month):
            jiangsu['url'].iloc[j]=url

jiangsu.drop('date_1',axis=1,inplace=True)
jiangsu.to_csv('./data/province/jiangsu/jiangsu.csv',encoding='gbk',index=False)