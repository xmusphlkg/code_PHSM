import os
from html import unescape
from urllib.parse import unquote
import numpy as np
import xlrd
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import requests


def remove_chinese(title):
    """
    The function `remove_chinese` takes a title as input and removes all non-digit and non-hyphen
    characters from it, returning the modified title and its type.
    
    :param title: The `title` parameter is a string that represents a file title
    :return: a tuple containing the modified title and the title type.
    """
    title_type = title.split('.')[-1]
    title = ''.join([x for x in title if x.isdigit() or x == '-'])
    return title, title_type


def filetype(title):
    """
    The function `filetype` takes a title as input and returns the file type by extracting the extension
    from the title.
    
    :param title: The title parameter is a string that represents the title of a file
    :return: The file type of the given title.
    """
    title_type = title.split('.')[-1]
    return title_type


def remove_space(title):
    """
    The function removes spaces from a given string.
    
    :param title: The parameter "title" is a string that represents a title or a sentence
    :return: the title with all spaces removed.
    """
    title = title.replace(' ', '')
    return title


def process_files_combined(directory):
    """
    Process files in the specified directory and extract data from different file formats (doc, docx, xls, xlsx, csv).
    
    Args:
        directory (str): The directory path where the files are located.
        
    Returns:
        pandas.DataFrame: A DataFrame containing the extracted data with columns ['疾病病种', '发病数', '死亡数', 'date'].
    """
    try:
        data_list = []
        datatype_list = (np.floating,np.integer,int, float)

        files = os.listdir(directory)

        for file in files:
            file_path = os.path.join(directory, file)
            file_size = os.path.getsize(file_path)
            print(file)

            if file_size > 100:
                file_type = filetype(file)
                if file_type == 'doc':
                    pass
                if file_type == 'docx':
                    text_content, table_content = read_docx(file_path)
                    df = pd.DataFrame(table_content)
                    for i in range(len(df)):
                        try:
                            df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                        except:
                            pass
                        if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                                df.iloc[i][2], datatype_list):
                            data_list.append(
                                [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
                elif file_type == 'xls':
                    sheet = xlrd.open_workbook(file_path).sheet_by_index(0)
                    num_rows = sheet.nrows
                    num_cols = sheet.ncols
                    data = []

                    for row_index in range(num_rows):
                        row_data = []

                        for col_index in range(num_cols):
                            cell_value = sheet.cell_value(row_index, col_index)
                            row_data.append(cell_value)

                        data.append(row_data)

                    df = pd.DataFrame(data)

                    for i in range(len(df)):
                        try:
                            df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                        except:
                            pass
                        if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                                df.iloc[i][2], datatype_list):
                            data_list.append(
                                [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
                elif file_type == 'csv':
                    df = pd.read_csv(file_path, encoding='gbk')
                    df.replace('-', '0', regex=True, inplace=True)
                    for i in range(len(df)):
                        if str(df.iloc[0, 0]) == '0':
                            try:
                                df.iloc[i, 2] = float(df.iloc[i, 2])
                                df.iloc[i, 3] = float(df.iloc[i, 3])
                            except:
                                pass
                            if isinstance(df.iloc[i][1], str) and isinstance(df.iloc[i][2], datatype_list) and isinstance(
                                    df.iloc[i][3], datatype_list):
                                data_list.append(
                                    [remove_space(df.iloc[i][1]), df.iloc[i][2], df.iloc[i][3], file.split('.')[0]])
                        else:
                            try:
                                df.iloc[i, 1] = float(df.iloc[i, 1])
                                df.iloc[i, 2] = float(df.iloc[i, 2])
                            except:
                                pass
                            try:
                                if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i,1],datatype_list) and isinstance(
                                        df.iloc[i][2], datatype_list):
                                    data_list.append(
                                        [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
                            except:
                                pass
                elif file_type == 'xlsx':
                    sheet = pd.read_excel(file_path)
                    df = sheet

                    for i in range(len(df)):
                        try:
                            df.iloc[i][1], df.iloc[i][2] = float(df.iloc[i][1]), float(df.iloc[i][2])
                        except:
                            pass
                        if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], datatype_list) and isinstance(
                                df.iloc[i][2], datatype_list):
                            data_list.append(
                                [remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
    except:
        pass

    result_df = pd.DataFrame(data_list, columns=['疾病病种', '发病数', '死亡数', 'date'])
    return result_df


def read_docx(file_path):
    """
    The function `read_docx` reads the content of a .docx file and returns the text content as a list of
    paragraphs and the table content as a list of rows.
    
    :param file_path: The file path is the location of the .docx file that you want to read. It should
    be a string that specifies the path to the file, including the file name and extension. For example,
    "C:/Documents/my_file.docx"
    :return: The function `read_docx` returns two values: `text_content` and `table_content`.
    """
    doc = Document(file_path)
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(row_content)
    return text_content, table_content


def get_year_month(title):
    """
    The function `get_year_month` takes a title as input and returns the year and month extracted from
    the title.
    
    :param title: The title of a text or document that contains information about a year and month
    :return: the year and month extracted from the given title.
    """
    year = title.split('布')[1].split('年')[0]
    month = title.split('年')[1].split('月')[0]
    return year, month


def download_and_parse_data(df_row, headers):
    """
    The function `download_and_parse_data` downloads and parses data from a given URL, saving it as a
    CSV file if it contains a table, and as a DOCX file if it contains a downloadable link.
    
    :param df_row: The `df_row` parameter is a row from a DataFrame. It contains information about a
    specific data entry, such as the URL, year, and month
    :param headers: The `headers` parameter is a dictionary that contains the headers to be sent with
    the HTTP request. These headers can include information such as user agent, content type, and
    authorization
    """
    url = df_row['链接']
    sign = 0

    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser', from_encoding='gbk')
    links_with_blank_target = soup.find_all('a', {'target': '_blank'})
    for link_tag in links_with_blank_target:
        link = link_tag.get('href')
        if link.endswith('docx'):
            sign = 1
            url_link = unescape(str(url) + link)
            response = requests.get(url_link, headers=headers)
            if response.status_code == 200:
                with open(f'./data/province/anhui/{df_row["年份"]}-{df_row["月份"]}.docx', 'wb') as f:
                    f.write(response.content)
            else:
                print(f'{df_row["年份"]}-{df_row["月份"]}存在，下载失败，但是200')

    try:
        table = soup.find('tbody')
        rows = table.find_all('tr')
        table_data = []
        for row in rows:
            cells = row.find_all(['td', 'th'])
            row_data = [cell.text.strip() for cell in cells]
            table_data.append(row_data)

        table_df = pd.DataFrame(table_data)
        table_df.replace('\xa0', '', regex=True, inplace=True)
        table_df.to_csv(f'./data/province/anhui/{df_row["年份"]}-{df_row["月份"]}.csv', encoding='gbk')
        sign = 1
    except Exception as e:
        print(f"Error processing table data: {e}")

    if sign == 0:
        print(f'{df_row["年份"]}-{df_row["月份"]}不存在传染病信息')


def scrape_and_save_table(soup, df, i):
    """
    The above code contains several functions related to scraping and manipulating data from tables,
    updating URL columns, modifying date formats, and finding missing elements in dataframes.
    
    :param soup: The parameter `soup` is a BeautifulSoup object that represents the HTML content of a
    webpage. It is used to extract data from the HTML structure
    :param df: The parameter `df` is a pandas DataFrame that contains the data you want to scrape and
    save
    :param i: The parameter "i" is used as an index to access specific columns or elements in the
    dataframe. It is used in functions like "scrape_and_save_table" and "find_missing_elements" to
    specify the column or element to be accessed
    """
    table = soup.find('tbody')
    rows = table.find_all('tr')
    table_data = []

    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_data = [cell.text.strip() for cell in cells]
        table_data.append(row_data)

    table_df = pd.DataFrame(table_data)
    table_df.replace('\xa0', '', regex=True, inplace=True)

    file_path = f'./data/province/anhui/{df["年份"].iloc[i]}-{df["月份"].iloc[i]}.csv'
    table_df.to_csv(file_path, encoding='gbk')


def update_url_column(df_csv_path, url_csv_path):
    """
    The above code contains three functions: "update_url_column" which updates a URL column in a
    DataFrame based on another DataFrame, "motify_date" which modifies the date format in a DataFrame,
    and "find_missing_elements" which finds missing elements in one column of a DataFrame compared to
    another column.
    
    :param df_csv_path: The file path of the CSV file containing the data frame
    :param url_csv_path: The `url_csv_path` parameter is the file path to the CSV file that contains the
    URLs
    """
    try:
        url_df = pd.read_csv(url_csv_path, encoding='gbk')
    except:
        url_df = pd.read_csv(url_csv_path)
    df = pd.read_csv(df_csv_path, encoding='gbk')
    try:
        df['date_1'] = pd.to_datetime(df['date'],format='%b-%y')
    except:
        df['date_1'] = pd.to_datetime(df['date'])
    df['year'] = df['date_1'].dt.year.astype(str)
    df['month'] = df['date_1'].dt.month.astype(str)
    df['url'] = None

    for i in range(len(url_df)):
        try:
            url = unquote(url_df.loc[i, '链接'])
        except:
            url = unquote(url_df.loc[i, 'url'])
        try:
            url_year = int(url_df['年份'].iloc[i])
        except:
            url_year = int(url_df['年'].iloc[i])
        try:
            url_month = int(url_df['月份'].iloc[i])
        except:
            url_month = int(url_df['月'].iloc[i])

        for j in range(len(df)):
            if str(df['year'].iloc[j]) == str(url_year) and str(df['month'].iloc[j]) == str(url_month):
                df['url'].iloc[j] = url
    df.drop('date_1', axis=1, inplace=True)
    df.to_csv(df_csv_path, encoding='gbk', index=False)

def motify_date(df,df_csv_path):
    """
    The function `motify_date` takes a DataFrame `df` and a file path `df_csv_path`, modifies the 'date'
    column in the DataFrame to a standardized format, drops any rows with missing values in the modified
    'date' column, adds separate '年' and '月' columns based on the modified 'date' column, and saves the
    modified DataFrame to a CSV file specified by `df_csv_path`.
    
    :param df: The parameter `df` is a pandas DataFrame that contains the data you want to modify
    :param df_csv_path: The path where you want to save the modified dataframe as a CSV file
    """
    df_1=df
    df_1['date'] = pd.to_datetime(df['date'], format='%b-%y', errors='coerce').dt.strftime('20%y-%m')
    df_1=df_1.dropna(axis=0)
    df_2=df
    df_2['date'] = pd.to_datetime(df['date'], format='%y-%b', errors='coerce').dt.strftime('20%y-%m')
    df_2=df_2.dropna(axis=0)
    df=pd.concat([df_1,df_2])
    df["年"]=df["date"].str.split("-",expand=True)[0]
    df["月"]=df["date"].str.split("-",expand=True)[1]
    df.to_csv(df_csv_path, encoding='gbk', index=False)

def find_missing_elements(df_1,i, df_2,m):
    str_list = df_1[~df_1[i].isin(df_2[m])][i].tolist()
    return str_list