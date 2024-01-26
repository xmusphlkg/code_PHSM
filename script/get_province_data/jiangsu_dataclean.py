import os
import pandas as pd
from docx import Document
from liu_script.zhejiang_dataclean import remove_space

files = os.listdir('./data/province/jiangsu')
def read_docx(file_path):
    doc = Document(file_path)

    # 逐段读取文本内容
    text_content = []
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)

    # 逐表格读取内容
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_content = [cell.text for cell in row.cells]
            table_content.append(row_content)

    return text_content, table_content

data_list = []
for file in files:
    text_content, table_content = read_docx('./data/province/jiangsu/'+file)
    df = pd.DataFrame(table_content)
    for i in range(len(df)):
        try:
            df.iloc[i][1],df.iloc[i][2] = float(df.iloc[i][1]),float(df.iloc[i][2])
        except:
            pass
        if isinstance(df.iloc[i][0], str) and isinstance(df.iloc[i][1], (int, float)) and isinstance(
                df.iloc[i][2], (int, float)):
            data_list.append([remove_space(df.iloc[i][0]), df.iloc[i][1], df.iloc[i][2], file.split('.')[0]])
result_df = pd.DataFrame(data_list, columns=['疾病病种', '发病数', '死亡数', 'date'])
result_df.to_csv('./data/province/jiangsu/jiangsu.csv', index=False, encoding='gbk')